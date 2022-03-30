from importlib.resources import path
from django.shortcuts import render
from flask import Flask,render_template,request,send_file,after_this_request

from docx import Document
import pandas as pd
import docx2txt
from io import BytesIO
import requests
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
from xml.etree import ElementTree
from io import StringIO
import io
import csv
import base64
import numpy as np
import zipfile
import os
from os.path import basename
from werkzeug.wrappers import Response

app=Flask(__name__)
UPLOAD_PATH = os.getcwd()
UPLOAD_PATH=UPLOAD_PATH+'\\'

@app.route('/',methods=['GET','POST'])
def index():
    return render_template('index.html')

@app.route('/database_download/<filename>')
def database_download(filename):
    #return send_from_directory('database_reports', filename)
    @after_this_request
    def remove_file(response): 
        print('Remove the file...In progress')
        #os.remove(UPLOAD_PATH+'Test.zip')

        return response 
    return send_file(filename, mimetype='zip', attachment_filename=filename, as_attachment=True)

@app.route('/data1',methods=['POST'])
def data1():
    data()
    #return send_file('Test.zip', mimetype='zip', attachment_filename='Test.zip', as_attachment=True)
    #return render_template('data.html')
    return render_template('data.html', filename='Test.zip')

@app.route('/data',methods=['POST'])
def data():
    if request.method=='POST':
        # f=request.form['tdd']
        
        # document = Document(f)
        # table=document.tables[1]
        # data = [[cell.text for cell in row.cells] for row in table.rows]

        # df = pd.DataFrame(data)
        # print(df)
        url = request.form['tdd']
        document = BytesIO(requests.get(url).content)
        document = Document(document)


        ##This function extracts the tables and paragraphs from the document object
        def iter_block_items(parent):
            """
            Yield each paragraph and table child within *parent*, in document order.
            Each returned value is an instance of either Table or Paragraph. *parent*
            would most commonly be a reference to a main Document object, but
            also works for a _Cell object, which itself can contain paragraphs and tables.
            """
            if isinstance(parent, doctwo):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("something's not right")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

            
        #This function extracts the table from the document object as a dataframe
        def read_docx_tables(tab_id=None, **kwargs):
            """
            parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)
            Parameters:
                filename:   file name of a Word Document
                tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                            When [None] - return a list of DataFrames (parse all tables)
                kwargs:     arguments to pass to `pd.read_csv()` function
            Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
            """
            def read_docx_tab(tab, **kwargs):
                vf = io.StringIO()
                writer = csv.writer(vf)
                for row in tab.rows:
                    writer.writerow(cell.text for cell in row.cells)
                vf.seek(0)
                return pd.read_csv(vf, **kwargs)

        #    doc = Document(filename)
            if tab_id is None:
                return [read_docx_tab(tab, **kwargs) for tab in document.tables]
            else:
                try:
                    return read_docx_tab(document.tables[tab_id], **kwargs)
                except IndexError:
                    print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                    raise



        #The combined_df dataframe will store all the content in document order including images, tables and paragraphs.
        #If the content is an image or a table, it has to be referenced from image_df for images and table_list for tables using the corresponding image or table id that is stored in combined_df
        #And if the content is paragraph, the paragraph text will be stored in combined_df
        combined_df = pd.DataFrame(columns=['para_text','table_id','style'])
        table_mod = pd.DataFrame(columns=['string_value','table_id'])

        #The image_df will consist of base64 encoded image data of all the images in the document
        image_df = pd.DataFrame(columns=['image_index','image_rID','image_filename','image_base64_string'])

        #The table_list is a list consisting of all the tables in the document
        table_list=[]
        xml_list=[]

        i=0
        imagecounter = 0


        blockxmlstring = ''
        for block in iter_block_items(document):
            if 'text' in str(block):
                isappend = False
                
                runboldtext = ''
                for run in block.runs:                        
                    if run.bold:
                        runboldtext = runboldtext + run.text
                        
                style = str(block.style.name)
        
                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n","")
                appendtxt = appendtxt.replace("\r","")
                tabid = 'Novalue'
                paragraph_split = appendtxt.lower().split()                
                
                isappend = True
                for run in block.runs:
                    xmlstr = str(run.element.xml)
                    my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                    root = ET.fromstring(xmlstr) 
                    #Check if pic is there in the xml of the element. If yes, then extract the image data
                    if 'pic:pic' in xmlstr:
                        xml_list.append(xmlstr)
                        for pic in root.findall('.//pic:pic', my_namespaces):
                            cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                            name_attr = cNvPr_elem.get("name")
                            blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                            embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            isappend = True
                            appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                            document_part = document.part
                            image_part = document_part.related_parts[embed_attr]
                            image_base64 = base64.b64encode(image_part._blob)
                            image_base64 = image_base64.decode()                            
                            dftemp = pd.DataFrame({'image_index':[imagecounter],'image_rID':[embed_attr],'image_filename':[name_attr],'image_base64_string':[image_base64]})
                            image_df = image_df.append(dftemp,sort=False)
                            style = 'Novalue'
                        imagecounter = imagecounter + 1
                    
            elif 'table' in str(block):
                isappend = True
                style = 'Novalue'
                appendtxt = str(block)
                tabid = i
                dfs = read_docx_tables(tab_id=i)
                dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[i],'style':[style]})
                table_mod = table_mod.append(dftemp,sort=False)
                table_list.append(dfs)
                i=i+1
            if isappend:
                    dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[tabid],'style':[style]})
                    combined_df=combined_df.append(dftemp,sort=False)
                    
        combined_df = combined_df.reset_index(drop=True)
        image_df = image_df.reset_index(drop=True)

        print(combined_df.index)
        #INBOUND UPLOAD PROCESS 1:  LOAD RESIDUALS
        archive_pipeline=[]
        filecheck_pipeline=[]
        validation_pipeline_final=[]
        load_source_pipeline=[]
        load_source_pipeline_final=[]
        p_file_checks_final=[]
        p_load_source_final=[]
        files=[]
        for ind in combined_df.index:
            #Table1 process details
            if(combined_df.iloc[ind]['para_text'].upper().startswith("INBOUND UPLOAD PROCESS 1")):
                table=document.tables[combined_df.iloc[ind+1]['table_id']]
                data = [[cell.text for cell in row.cells] for row in table.rows]

                df = pd.DataFrame(data)
                df.columns=df.iloc[0]
                df = df[1:]
                process_name=df["Process Name"][1].strip()
                process_type=df["Process Type"][1].strip()
                process_short_name=df["Process Short Name"][1].strip()
                upload_batch_size=df["Upload Batch Size"][1].strip()
                calendar_year=df["Calendar Year"][1].strip()
                start_period=df["Start Period"][1].strip()
                end_period=df["End Period"][1].strip()
                
                    #Table2 File/Source details
            if(combined_df.iloc[ind]['para_text'].upper()=="SOURCE DATA:"):

                src_det=document.tables[combined_df.iloc[ind+2]['table_id']]
                data = [[cell.text for cell in row.cells] for row in src_det.rows]
                    
                src_dict={}
                num_of_src=len(data)-1
                while num_of_src>0:
                    src_dict[num_of_src]=data[num_of_src]
                    num_of_src-=1
                    #print(src_dict)
                        
                    src_list=[]
                for k,v in reversed(src_dict.items()):
                    
                            #v[0] sourcetype
                            #v[1] label
                            #v[2] file/object name
                            #v[3] location
                            #v[4] properties
                            #v[5] category( source/xref)
                            #v[6] frequency
                            #v[7] volume
                            #v[8] abort on load error
                            #v[9] archive
                            #v[10] encrpted
                            #v[11] decryption key
                        #load_source_pipeline=[]    
                    if v[0].upper()=='FILE':
                        src_list.append(f'''create step if not exists s_set_source_file_name_{process_short_name}_{v[1]} as (set v_filename_{process_short_name}_{v[1]} *= select name  from  DirList(Directory='{v[3]}', Filter='{v[2]}' )); ''')
                            
                        src_list.append(f'''create step if not exists s_set_abortflag_{process_short_name}_{v[1]} as (set v_abort_ondata_error_{process_short_name}_{v[1]} *= true);''')
                    
                        src_list.append(f'''create step if not exists s_load_clean_{process_short_name}_{v[1]} as (insert into Delta( TableName= 'delta.{process_short_name}_{v[1]}_clean' , Unlogged=true, Overwrite=true ) select * from delta.{process_short_name}_{v[1]}_dump where (:v_filename_{process_short_name}_{v[1]} ||'_'||row_no) not in (select distinct source_file||'_'||row_no from delta.validation_errors_{process_short_name}));''')
                    
                        src_list.append(f'''create step if not exists s_set_fileread_error_email_subj_{process_short_name}_{v[1]} as (set v_email_subject *=:v_shared_customer_name ||'Process '||:v_process_name_{process_short_name}||' could not be started.Field name mismatch for file '|| :v_filename_{process_short_name}_{v[1]} );''')
                    
                        src_list.append(f'''create step if not exists s_send_fileread_error_email_{process_short_name}_{v[1]} as (send email e_generic_email);''')
                    
                        src_list.append(f'''create step if not exists s_set_nofile_email_subj_{process_short_name}_{v[1]} as (set v_email_subject *=:v_shared_customer_name ||'- '||:v_process_name_{process_short_name}||' process could not be started. File '||'{v[2]}'||' not found for period '|| :v_period_name );''')
                    
                        src_list.append(f'''create step if not exists s_set_multiplefiles_email_subj_{process_short_name}_{v[1]} as (set v_email_subject *=:v_shared_customer_name ||'- '||:v_process_name_{process_short_name}||' process could not be started. Multiple  '||'{v[2]}'||' Files  found for period '|| :v_period_name);''')
                    
                        src_list.append(f'''create step if not exists s_send_nofile_email_{process_short_name}_{v[1]} as (send email e_generic_email);''')
                    
                        src_list.append(f'''create step if not exists s_send_multiplefiles_email_{process_short_name}_{v[1]} as (send email e_generic_email);''')
                        
                        #FILECHECK PIPELINES/READ
                        filecheck_pipeline.append(f'''create pipeline if not exists p_filechecks_{process_short_name}_{v[1]} (ContinueOnError=false);''')

                        filecheck_pipeline.append(f'''alter pipeline p_filechecks_{process_short_name}_{v[1]} add step s_set_nofile_email_subj_{process_short_name}_{v[1]};''')

                        filecheck_pipeline.append(f'''alter pipeline p_filechecks_{process_short_name}_{v[1]} add step s_shared_no_action as (Condition=(select IF count(name) =0 THEN false ELSE true END  from  DirList(Directory='{v[3]}', Filter='{v[2]}' )  ), OnConditionFalse=s_send_nofile_email_{process_short_name}_{v[1]}, AbortOnConditionFalse=true);''')

                        filecheck_pipeline.append(f'''alter pipeline p_filechecks_{process_short_name}_{v[1]} add step s_set_multiplefiles_email_subj_{process_short_name}_{v[1]};''')

                        filecheck_pipeline.append(f'''alter pipeline p_filechecks_{process_short_name}_{v[1]} add step s_shared_no_action as (Condition=(select IF count(name) >1 THEN false ELSE true END  from  DirList(Directory='{v[3]}', Filter='{v[2]}' ) ), OnConditionFalse=s_send_multiplefiles_email_{process_short_name}_{v[1]}, AbortOnConditionFalse=true);''')
                        
                        filecheck_pipeline.append('\n')
                        
                        
                        filecheck_pipeline.append(f'''create pipeline if not exists p_report_fileread_error_{process_short_name}_{v[1]} (ContinueOnError=false);''')

                        filecheck_pipeline.append(f'''alter pipeline p_report_fileread_error_{process_short_name}_{v[1]} add step s_set_fileread_error_email_subj_{process_short_name}_{v[1]};''')

                        filecheck_pipeline.append(f'''alter pipeline p_report_fileread_error_{process_short_name}_{v[1]} add step s_send_fileread_error_email_{process_short_name}_{v[1]};''')
                        
                        filecheck_pipeline.append('\n')
                        #READ
                        filecheck_pipeline.append(f'''create pipeline if not exists p_read_source_file_{process_short_name}_{v[1]} (ContinueOnError=false, OnError=p_report_fileread_error_{process_short_name}_{v[1]});''')


                        filecheck_pipeline.append(f'''alter pipeline p_read_source_file_{process_short_name}_{v[1]} add step s_set_source_file_name_{process_short_name}_{v[1]};''')


                        filecheck_pipeline.append(f'''alter pipeline p_read_source_file_{process_short_name}_{v[1]} add step s_read_source_dump_{process_short_name}_{v[1]};''')

                        filecheck_pipeline.append('\n')
                        
                        #aggregation FILECHECk/LOAD SOURCE
                        p_file_checks_final.append(f'''alter pipeline p_kickoff_{process_short_name} add pipeline p_filechecks_{process_short_name}_{v[1]};''')
                        p_load_source_final.append(f'''alter pipeline p_kickoff_{process_short_name} add pipeline p_load_source_file_{process_short_name}_{v[1]};''')

                        
                        #LOAD SOURCE PIPELINE
                        
                        load_source_pipeline.append(f'''create pipeline if not exists p_load_source_file_{process_short_name}_{v[1]} (ContinueOnError=false);''')


                        load_source_pipeline.append(f'''alter pipeline p_load_source_file_{process_short_name}_{v[1]} add pipeline p_read_source_file_{process_short_name}_{v[1]};''')


                        load_source_pipeline.append(f'''alter pipeline p_load_source_file_{process_short_name}_{v[1]} add pipeline p_validate_source_file_{process_short_name}_{v[1]};''')
                        load_source_pipeline.append('\n')
                        load_source_pipeline_final.append(f'''alter pipeline p_load_sourcedata_{process_short_name} add pipeline p_load_source_file_{process_short_name}_{v[1]};''')

                        #ARCHIVE
                        if(v[9]=='True'):
                            if not archive_pipeline:
                                archive_pipeline.append(f'''create pipeline if not exists p_archive_sourcefiles_{process_short_name} (ContinueOnError=false);''')
                    
                            archive_pipeline.append(f'''alter pipeline p_archive_sourcefiles_{process_short_name} add step s_move_to_archive_{process_short_name}_{v[1]};''')
                            archive_pipeline.append(f'''alter pipeline p_archive_sourcefiles_{process_short_name} add step s_add_archive_file_time_{process_short_name}_{v[1]};''')



                            src_list.append(f'''create step if not exists s_move_to_archive_{process_short_name}_{v[1]} as (call MoveFile(FilePath='{v[3]}'||:v_filename_{process_short_name}_{v[1]}, Destination='{v[3]}archive/'));''')
                    
                            src_list.append(f'''create step if not exists s_add_archive_file_time_{process_short_name}_{v[1]} as (call RenameFile(FilePath='{v[3]}archive/'||:v_filename_{process_short_name}_{v[1]}, Name='/incoming/inbound/archive/'||:v_filename_{process_short_name}_{v[1]}||'_'||FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss')));''')
                    
                        src_list.append('\n')
                
                        for ind1 in combined_df.index:
                            
                        #Source Layout specification
                            if(combined_df.iloc[ind1]['para_text'].upper().strip()==(f'''SOURCE LAYOUT: {v[1]}''')):
                            
                                src_layout=document.tables[combined_df.iloc[ind1+1]['table_id']]
                                data1 = [[cell.text for cell in row.cells] for row in src_layout.rows]
                                df_layout = pd.DataFrame(data1)

                                df_layout.columns=df_layout.iloc[0]
                                df_layout = df_layout[1:]
                                mandatory_columns=df_layout.loc[df_layout['Mandatory'] =='YES']
                                #dateformat_columns= df_layout.loc[df_layout['Data Type'].str.upper() != 'DATE']
                                datatype_columns= df_layout.loc[(df_layout['Data Type'].str.upper() =='NUMBER')|(df_layout['Data Type'].str.upper() =='DATE')]
                                #print(datatype_columns)
                                #final validations to be appended
                                validations_final=[]
                                data_type_validation=[]
                                mandatory_check_validation=[]
                                all_src_fields=[]
                                validation_pipeline=[]
                                
                                if not validation_pipeline:
                                    validation_pipeline.append('\n')
                                    validation_pipeline.append(f'''create pipeline if not exists p_validate_source_file_{process_short_name}_{v[1]} (ContinueOnError=false);''')
                                    validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_validate_source_{process_short_name}_{v[1]};''')




                            
                                #for getting all columns
                                
                                for index, row in df_layout.iterrows():
                                    src_layout_all_field_name=row["Field Name"].replace(" ","_") 
                                    if src_layout_all_field_name:
                                        all_src_fields.append(f'''"{src_layout_all_field_name}" as {src_layout_all_field_name.replace(" ","_")}''')
                                
                                validations_final.append((f'''create step if not exists s_read_source_dump_{process_short_name}_{v[1]} as (insert into Delta( TableName= 'delta.{process_short_name}_{v[1]}_dump' , Unlogged=true, Overwrite=true ) select ''')+(', '.join(all_src_fields))+(f''',SeqNum() as row_no from  ReadFile(FilePath='{v[3]}'||:v_filename_{process_short_name}_{v[1]} , FirstLineNames=true, Separator=',', Quote='"', Trim=true) );'''))
                                
                                
                                #for mandatory columns
                                for index, row in mandatory_columns.iterrows():
                                    src_layout_mandatory_field_name=row["Field Name"].replace(" ","_")
                                    src_layout_mandatory_data_type=row["Data Type"]
                                    src_layout_mandatory=row["Mandatory"]
                                    if src_layout_mandatory.upper() =='YES':
                                        validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_validate_mandatory_{process_short_name}_{v[1]}_{src_layout_mandatory_field_name.replace(" ","_")};''')
                                        validations_final.append((f'''Create step if not exists s_validate_mandatory_{process_short_name}_{v[1]}_{src_layout_mandatory_field_name} as (insert into Delta( TableName= 'delta.validate_{process_short_name}_{v[1]}' , Unlogged=true, Overwrite=true ) select :v_filename_{process_short_name}_{v[1]},row_no,'{src_layout_mandatory_field_name}', Nvl({src_layout_mandatory_field_name},NULL),'Manadatory Data missing' from delta.{process_short_name}_{v[1]}_dump where {src_layout_mandatory_field_name.replace(" ","_")} is null);'''))
                                
                                #for datatype columns    
                                for index, row in datatype_columns.iterrows():
                                    src_layout_field_name=row["Field Name"].replace(" ","_") 
                                    src_layout_data_type=row["Data Type"]
                                    src_layout_format=row["Format"]
                                    if src_layout_data_type.upper() =='DATE':
                                        validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_validate_dataType_format_{process_short_name}_{v[1]}_{src_layout_field_name.replace(" ","_")};''')
                                        validations_final.append((f'''create step if not exists s_validate_dataType_format_{process_short_name}_{v[1]}_{src_layout_field_name.replace(" ","_")} as (insert into Delta( TableName= 'delta.validation_errors_{v[1]}'  ) select :v_filename_{process_short_name}_{v[1]} ,row_no,'{src_layout_field_name}',Nvl({src_layout_field_name},'NULL'),'Invalid date {src_layout_format}' from delta.validate_{process_short_name}_{v[1]} where {src_layout_field_name} is not null and isvalid_{src_layout_field_name}=false);'''))
                                        data_type_validation.append(f'''{src_layout_field_name.replace(" ","_")}, IsDateTimeFormat({src_layout_field_name.replace(" ","_")},'{src_layout_format}') isvalid_{src_layout_field_name.replace(" ","_")}''')
                                    else:
                                        validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_validate_dataType_format_{process_short_name}_{v[1]}_{src_layout_field_name.replace(" ","_")};''')
                                        validations_final.append((f'''create step if not exists s_validate_dataType_format_{process_short_name}_{v[1]}_{src_layout_field_name.replace(" ","_")} as (insert into Delta( TableName= 'delta.validation_errors_{v[1]}'  ) select :v_filename_{process_short_name}_{v[1]} ,row_no,'{src_layout_field_name.replace(" ","_")}',Nvl({src_layout_field_name.replace(" ","_")},'NULL'),'Invalid decimal ' from delta.validate_{process_short_name}_{v[1]} where {src_layout_field_name.replace(" ","_")} is not null and isvalid_{src_layout_field_name.replace(" ","_")}=false);'''))
                                        data_type_validation.append(f'''{src_layout_field_name.replace(" ","_")}, IsNumberFormat({src_layout_field_name.replace(" ","_")},'###.##')''' )
                                if len(validations_final)>1:
                                    validations_final.append((f'''Create step if not exists s_validate_source_{process_short_name}_{v[1]} as (insert into Delta( TableName= 'delta.validate_{process_short_name}_{v[1]}' , Unlogged=true, Overwrite=true ) select :v_filename_{process_short_name}_{v[1]}, row_no,''')+(', '.join(data_type_validation))+(f''' from delta.{process_short_name}_{v[1]}_dump);'''))
                                    validations_final.append('\n')
                                    src_list.append('\n')
                                    src_list.extend(validations_final)
                                    validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_set_abortflag_{process_short_name}_{v[1]};''')
                                    validation_pipeline.append(f'''alter pipeline p_validate_source_file_{process_short_name}_{v[1]} add step s_load_clean_{process_short_name}_{v[1]} as (Condition=select if count(*)=0 or  :v_abort_ondata_error_{process_short_name}_{v[1]}=false  then true else false end from delta.validation_errors_{process_short_name} where source_file=:v_filename_{process_short_name}_{v[1]}, OnConditionFalse=p_report_validation_errors_{process_short_name}, AbortOnConditionFalse=true);''')

                                    validation_pipeline_final.extend(validation_pipeline)

        #create step if not exists s_set_processname_RES as (set v_process_name_RES *= 'Load  Residuals' );
        steps=[];
        steps.append(f'''create step if not exists s_set_start_email_subject_{process_short_name} as (set v_email_subject *=:v_shared_customer_name ||' - Process '||:v_process_name_{process_short_name}||' started for period '||:v_period_name);''')

        steps.append(f'''create step if not exists s_send_start_email_{process_short_name} as (send email e_generic_email);''')

        steps.append(f'''create step if not exists s_send_finalize_email_{process_short_name} as (send email e_generic_email); ''')

        steps.append(f'''create step if not exists s_set_batch_size_{process_short_name} as (set v_batch_size *= {upload_batch_size});''')

        steps.append(f'''create step if not exists s_set_processname_{process_short_name} as (set v_process_name_{process_short_name} *='{process_name}');''')

        steps.append(f'''create step if not exists s_set_pipelinename_{process_short_name} as (set v_pipeline_name *= 'p_upload_process_{process_short_name}' );''')

        #can be modified as per error boundries
        steps.append(f'''create step if not exists s_set_e_BG_finalized_subject_{process_short_name} as (set v_email_subject *=:v_shared_customer_name ||'- '||:v_process_name_{process_short_name}||' process could not be started. One or more Business Groups for the period '|| :v_period_name ||' are already finalized.');''')

        steps.append(f'''create step if not exists s_create_validation_errorlog_{process_short_name} as (insert into Delta( TableName='delta.validation_errors_{process_short_name}' , Unlogged=true, Overwrite=true ) select '' as source_file, '' as row_no, '' as data_field, '' as field_value, '' as error_message from Empty() where '1'='2');''')

        steps.append(f'''create step if not exists s_set_source_validation_log_filename_{process_short_name} as (set v_sourcevalidation_error_log_{process_short_name} *= :v_period_name ||'_sourcedatavalidation_errors_{process_short_name}' || FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss') ||'.csv');''')

        steps.append(f'''create step if not exists s_generate_sourcevalidations_error_log_{process_short_name} as (call  WriteFile(FilePath='/logs/' ||:v_sourcevalidation_error_log_{process_short_name},Input= (select * from delta.validation_errors_{process_short_name} ), FirstLineNames=true,Separator=',',Quote='"',Trim=true));''')

        steps.append(f'''create step if not exists s_create_sourcevalidation_error_email_{process_short_name} as (create email if not exists e_sourcevalidation_error_email_{process_short_name} as ("From"=:v_email_from,"To"=:v_email_to, "Body"=:v_generic_email_body, "BodyType"='html', "Subject"=:v_shared_customer_name ||'Process '||:v_process_name_{process_short_name}||' could not be started.Field validations failed', "Attachment1" = ( '/logs/' || :v_sourcevalidation_error_log_{process_short_name}),"AttachmentType1" = 'plain' ));''')

        steps.append(f'''create step if not exists s_send_sourcevalidation_error_email_{process_short_name} as (send email e_sourcevalidation_error_email_{process_short_name});''')


        steps.append(f'''create step if not exists s_set_process_log_file_name_{process_short_name} as (set v_process_log_file_name_{process_short_name} *=:v_period_name || '_{process_short_name}_Rejects_' || FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss') ||'.csv'); ''')

        steps.append(f'''create step if not exists s_set_incenterr_log_file_name_{process_short_name} as (set v_incenterr_log_file_name_{process_short_name} *=:v_period_name || '_{process_short_name}_Incent_Error_Log_' || FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss') ||'.csv');''')

        steps.append(f'''create step if not exists s_set_logs_zip_file_name_{process_short_name} as (set v_logs_zip_file_name_{process_short_name} *=:v_period_name || '_{process_short_name}_Logs_' || FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss') ||'.zip');''')

        steps.append(f'''create step if not exists s_create_process_log_file_{process_short_name} as (call  WriteFile(FilePath='/logs/temp/' ||:v_process_log_file_name_{process_short_name},Input= (select * from delta.process_log), FirstLineNames=true,Separator=',',Quote='"',Trim=true));''')

        steps.append(f'''create step if not exists s_create_incent_error_log_file_{process_short_name} as (call  WriteFile(FilePath='/logs/temp/' ||:v_incenterr_log_file_name_{process_short_name},Input= (select * from delta.archive_order_item_validation_error ), FirstLineNames=true,Separator=',',Quote='"',Trim=true));''')

        steps.append(f'''create step if not exists s_zip_all_log_files_{process_short_name} as (call ZipFile( Directory= '/logs/temp', Destination= '/logs/'||:v_logs_zip_file_name_{process_short_name}, Overwrite=true));''')

        steps.append(f'''create step if not exists s_create_process_completed_email_with_attachment_{process_short_name} as (create email if not exists e_process_completed_email_with_att_{process_short_name} as ("From"=:v_email_from,"To"=:v_email_to, "Body"=:v_generic_email_body, "BodyType"='html', "Subject"=:v_shared_customer_name ||' - Process '||:v_process_name_{process_short_name}||' has COMPLETED for period '||:v_period_name, "Attachment1" = ( '/logs/' || :v_logs_zip_file_name_{process_short_name}),"AttachmentType1" = 'zip' ));''')

        steps.append(f'''create step if not exists s_create_process_completed_email_no_attachment_{process_short_name} as (create email if not exists e_process_completed_email_no_att_{process_short_name} as ("From"=:v_email_from,"To"=:v_email_to, "Body"='<b><font color="red">Note : Due to huge log size, logs have been copied to FTP server at "/logs/' || :v_logs_zip_file_name_{process_short_name}||'" </font></b><br/>'||:v_generic_email_body, "BodyType"='html', "Subject"=:v_shared_customer_name ||' - Process '||:v_process_name_{process_short_name}||' has COMPLETED for period '||:v_period_name ));''')

        steps.append(f'''create step if not exists s_send_process_completed_email_with_att_{process_short_name} as (send email e_process_completed_email_with_att_{process_short_name});''')

        steps.append(f'''create step if not exists s_send_process_completed_email_no_att_{process_short_name} as (send email e_process_completed_email_no_att_{process_short_name});''')

        steps.append(f'''create step if not exists s_update_delta_pgqueue_mapping_with_period_{process_short_name} as (Insert into Delta(TableName='delta.pgqueue_mapping_with_period',Overwrite=true, Unlogged=true) select parent_process_group_name,child_process_group_name,process_group_order,Uppercase(param_name) as param_name,param_value ,is_variable, :v_period_name  as period_name from pgconfig.pgqueue_mapping pgqm where pgqm.parent_process_group_name=:v_process_name);''')

        #for step in steps:
            #print(step,steps.index(step))
        steps.append('\n')
        steps.extend(src_list)
        #dfs = pd.DataFrame(steps)


        pipelines=[];

        pipelines.append('\n')
        pipelines.append('''--Pipeline Build starts here''')
        pipelines.append('\n')
        pipelines.append(f'''create pipeline if not exists p_custom_order_validations_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''--start:  pipeline to send email with all data type validations issues in all the source files. To report any errors from SFDC and Incent source, Make sure the data is copied to table  'delta.validation_errors_{process_short_name}''') 
        pipelines.append('\n')
        pipelines.append(f'''create pipeline if not exists p_report_validation_errors_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''alter pipeline p_report_validation_errors_{process_short_name} add step s_set_source_validation_log_filename_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_report_validation_errors_{process_short_name} add step s_generate_sourcevalidations_error_log_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_report_validation_errors_{process_short_name} add step s_create_sourcevalidation_error_email_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_report_validation_errors_{process_short_name} add step s_send_sourcevalidation_error_email_{process_short_name};''')

        pipelines.append(f'''--end :   pipeline to send email with all data type validations issues in all the source files. To report any errors from SFDC and Incent source, Make sure the data is copied to table  'delta.validation_errors_{process_short_name}' ''')

        pipelines.append('\n')

        pipelines.append(f'''create pipeline if not exists p_transform_{process_short_name} (ContinueOnError=false);''')



        pipelines.append(f'''--start: finalization check for upload orders {process_name} ''')

        pipelines.append(f'''create pipeline if not exists p_finalize_check_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''alter pipeline p_finalize_check_{process_short_name} add step s_set_e_BG_finalized_subject_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_finalize_check_{process_short_name} add step s_shared_no_action as (Condition=(select IF count(bg.name)=0 then true else false end    from xactly.xc_business_group bg join xactly.xc_finalize_business_group fb   on bg.business_group_id = fb.business_group_id    join xactly.xc_period pr   on  pr.period_id = fb.period_id AND pr.NAME = :v_period_name    join xactly.xc_finalize_payment fp    on fp.finalize_id=fb.finalize_id AND fp.finalize_type='Finalize'  ), OnConditionFalse=s_send_finalize_email_{process_short_name}, AbortOnConditionFalse=true);''')

        pipelines.append(f'''--end :  finalization check for upload orders {process_name}''') 

        #ARCHIVE
        pipelines.append('\n')
        pipelines.extend(archive_pipeline)

        #FILECHECK/READ
        pipelines.append('\n')
        pipelines.extend(filecheck_pipeline)

        #VALIDATION PIPELINE
        pipelines.append('\n')
        pipelines.extend(validation_pipeline_final)

        #LOAD SOURCE PIPELINE
        pipelines.append('\n')
        pipelines.extend(load_source_pipeline) 

        #STANDARD/CUSTOM VALIDATIONS
        pipelines.append('\n')
        pipelines.append(f'''create pipeline if not exists p_order_validations_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''alter pipeline p_order_validations_{process_short_name} add pipeline p_custom_order_validations_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_order_validations_{process_short_name} add pipeline p_standard_order_validations;''')


        pipelines.append(f'''alter pipeline p_order_validations_{process_short_name} add step s_insert_into_staging_valid_prestage_order_assignments;''')


        pipelines.append(f'''alter pipeline p_order_validations_{process_short_name} add step s_insert_into_staging_valid_prestage_order;''')

        #LOAD SOURCE DATA PIPELINE FINAL 
        pipelines.append(f'''create pipeline if not exists p_load_sourcedata_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''alter pipeline p_load_sourcedata_{process_short_name} add step s_create_validation_errorlog_{process_short_name};''')

        pipelines.extend(load_source_pipeline_final)


        #KICKOFF PROCESS BEGINS
        pipelines.append('\n')
        pipelines.append('''--KICKOFF PROCESS BEGINS''')
        pipelines.append(f'''create pipeline if not exists p_kickoff_{process_short_name} (ContinueOnError=false);''')

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add pipeline p_set_dynamic_variables;''')

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add step s_set_processname_{process_short_name};''')

        pipelines.extend(p_file_checks_final)

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add step s_create_validation_errorlog_{process_short_name};''')

        pipelines.extend(p_load_source_final)

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add pipeline p_finalize_check_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add step s_update_delta_pgqueue_mapping_with_period_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_kickoff_{process_short_name} add pipeline p_shared_queue_pg;''')
        pipelines.append('''--KICKOFF PROCESS ENDS''')
        #KICKOFF PROCESS ENDS

        #PREPROCESS BEGINS
        pipelines.append('\n')
        pipelines.append('''--PREPROCESS BEGINS''')
        pipelines.append(f'''create pipeline if not exists p_pre_process_{process_short_name} (ContinueOnError=false);''')


        pipelines.append(f'''alter pipeline p_pre_process_{process_short_name} add pipeline p_set_dynamic_variables;''')


        pipelines.append(f'''alter pipeline p_pre_process_{process_short_name} add step s_set_start_email_subject_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_pre_process_{process_short_name} add step s_send_start_email_{process_short_name};''')
        pipelines.append('''--PREPROCESS ENDS''')
        #PREPROCESS ENDS

        #MAIN UPLOAD PIPELINE BEGINS
        pipelines.append('\n')
        pipelines.append('''--MAIN UPLOAD PIPELINE BEGINS''')
        pipelines.append(f'''create pipeline if not exists p_upload_process_{process_short_name} (ContinueOnError=false, OnError=p_email_invocation_errors);''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_set_dynamic_variables;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_set_processname_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_set_pipelinename_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_set_batch_size_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_load_sourcedata_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_create_prestage_order_item;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_create_prestage_order_item_assignment;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_shared_create_archive_error_log;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_shared_create_process_log;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_transform_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_shared_delete_staging_tables;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add step s_set_order_item_field_list;''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_order_validations_{process_short_name};''')


        pipelines.append(f'''alter pipeline p_upload_process_{process_short_name} add pipeline p_shared_upload_orders as (Condition=select if count(name) =0 then true else false end from (show variables) where name='v_debug_flag_{process_short_name}' and value=true, OnConditionFalse=s_shared_no_action, AbortOnConditionFalse=false);''')
        pipelines.append('''--MAIN UPLOAD PIPELINE ENDS''')
        #MAIN UPLOAD PIPELINE ENDS


        #POSTPROCESS BEGINS

        pipelines.append('\n')
        pipelines.append('''--POSTPROCESS BEGINS''')
        pipelines.append(f'''create pipeline if not exists p_post_process_{process_short_name} (ContinueOnError=false);''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_create_temp_dir;''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_set_process_log_file_name_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_set_incenterr_log_file_name_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_set_logs_zip_file_name_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_create_incent_error_log_file_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_create_process_log_file_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_set_source_validation_log_filename_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_generate_sourcevalidations_error_log_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_copy_sourcevalidations_error_log_temp_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_zip_all_log_files_{process_short_name};''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_create_process_completed_email_with_attachment_{process_short_name} as (Condition=(select if size>5000000 then false else true end from DirList(Directory='/logs/', filter=:v_logs_zip_file_name_{process_short_name}) ), OnConditionFalse=s_create_process_completed_email_no_attachment_{process_short_name}, AbortOnConditionFalse=false);''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_send_process_completed_email_with_att_{process_short_name} as (Condition=(select if size>5000000 then false else true end from DirList(Directory='/logs/', filter=:v_logs_zip_file_name_{process_short_name}) ), OnConditionFalse=s_send_process_completed_email_no_att_{process_short_name}, AbortOnConditionFalse=false);''')

        pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add step s_delete_temp_dir;''')

        if len(archive_pipeline)>1:
            
            pipelines.append(f'''alter pipeline p_post_process_{process_short_name} add pipeline p_archive_sourcefiles_{process_short_name};''')
        pipelines.append('''--POSTPROCESS ENDS''')
        #POSTPROCESS ENDS

        steps.extend(pipelines)

        #WRITES ALL STEPS AND PIPELINE TO FILE

        np.savetxt(f'''{process_short_name}_connect_steps.txt''',steps, fmt='%s', delimiter="\t")
        files.append(f'''{process_short_name}_connect_steps.txt''')

        print('Completed!!!!!!!!!!!!!!!')


        #COMMON STEPS AND PIPELINES GENERATION
        common=[]

        common.append('''create step if not exists s_set_v_email_to as (set v_email_to *=:Email_Distribution_List);''')

        common.append('''create step if not exists s_period_start_date as (set v_period_start_date *= select start_date from xactly.xc_period where name=:v_period_name);''')

        common.append('''create step if not exists s_period_end_date as (set v_period_end_date *= select end_date from xactly.xc_period where name=:v_period_name);''')

        common.append('''create step if not exists s_shared_no_action;''')

        common.append('''create step if not exists s_set_generic_email_body as (set v_generic_email_body *= 'Auto generated by Xactly Connect <br> For documentation: https://community.xactlycorp.com/ <br>For help, email Xactly Support: support@xactlycorp.com');''')

        common.append('''create step if not exists s_set_generic_email_subject as (set v_email_subject *=:v_shared_customer_name ||'- Processes could not be started for the period '||:v_period_name);''')

        common.append('''create step if not exists s_create_generic_email as (create email if not exists e_generic_email as ("From"=:v_email_from,"To"=:v_email_to, "Body"=:v_generic_email_body, "BodyType"='html', "Subject"=:v_email_subject));''')

        common.append('''create step if not exists s_delete_stage_order_item as (delete from staging.order_item);''')

        common.append('''create step if not exists s_delete_stage_order_item_asgnmt as (delete from staging.order_item_assignment);''')

        common.append('''create step if not exists s_delete_stage_order_item_val_err as (delete from staging.order_item_validation_error);''')

        common.append('''create step if not exists s_delete_staging_geography as (delete from staging.geography);''')

        common.append('''create step if not exists s_delete_staging_geography_exception as (delete from staging.geography_exception);''')

        common.append('''create step if not exists s_delete_staging_customer as (delete from staging.customer);''')

        common.append('''create step if not exists s_delete_staging_customer_exception as (delete from staging.customer_exception);''')

        common.append('''create step if not exists s_delete_staging_product as (delete from staging.product);''')

        common.append('''create step if not exists s_delete_staging_product_exception as (delete from staging.product_exception);''')

        common.append('''create step if not exists s_shared_create_archive_error_log as (insert into Delta(TableName='delta.archive_order_item_validation_error', Overwrite=true , Unlogged=true) select * from staging.order_item_validation_error where '1'='2');''')

        common.append('''create step if not exists s_shared_create_process_log as (Insert into Delta(TableName='delta.process_log', Unlogged=true, Overwrite=true)  select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'' as category,'' as order_code,'' as item_code, '' as error_field, '' AS reject_reason from Empty() where '1'='2');''')

        common.append('''create step if not exists s_insert_into_staging_valid_prestage_order as (insert into staging.order_item  ({:v_order_item_field_list},batch_name) select {:v_order_item_field_list},batch_name||'_001' as batch_name from delta.prestage_order_item  where order_code||item_code not inn (select distinct order_code||item_code from delta.process_log where category='REJECT' ));''')

        common.append('''create step if not exists s_insert_into_staging_valid_prestage_order_assignments as (insert into staging.order_item_assignment (order_code,item_code,employee_id,split_amount_pct) select order_code, item_code, employee_id, split_amount_pct from delta.prestage_order_item_assignment where order_code||item_code not inn (select distinct order_code||item_code from delta.process_log where category='REJECT' ));''')

        common.append('''create step if not exists s_order_validate_mandatory_fields as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, '' as error_field, 'Order_Code or Item code is null' AS reject_reason from delta.prestage_order_item where order_code is null or item_code is null union select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, amount as error_field, 'Order Amount is null' AS reject_reason from delta.prestage_order_item where amount is null  union select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, amount_unit_type_name as error_field, 'Amount Unit type is  null' AS reject_reason from delta.prestage_order_item where amount_unit_type_name is null  union select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, incentive_date as error_field, 'Incentive Date is null' AS reject_reason from delta.prestage_order_item where incentive_date is null union select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, period_name as error_field, 'Period Name is null' AS reject_reason from delta.prestage_order_item where period_name is null  union select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, batch_name as error_field, 'Batch Name is null' AS reject_reason from delta.prestage_order_item where batch_name is null);''')

        common.append('''create step if not exists s_order_validate_batch_type as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, batch_type_name as error_field, 'Batch type is  invalid' AS reject_reason from delta.prestage_order_item where batch_type_name is not null and batch_type_name not in (select distinct NAME from xactly.xc_batch_type ));''')

        common.append('''create step if not exists s_order_validate_amount_unit_type as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, amount_unit_type_name as error_field, 'Amount Unit type is  invalid' AS reject_reason from delta.prestage_order_item where amount_unit_type_name is not null and amount_unit_type_name not in (select distinct NAME from xactly.xc_unit_type ));''')

        common.append('''create step if not exists s_order_validate_order_type as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, order_type_name as error_field, 'Order  type is  invalid' AS reject_reason from delta.prestage_order_item where order_type_name is not null and order_type_name not in (select distinct NAME from xactly.xc_order_type ));''')

        common.append('''create step if not exists s_order_validate_employee_id as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, employee_id as error_field, 'Invalid Employee Id' AS reject_reason from delta.prestage_order_item_assignment where employee_id not in (select distinct employee_id from xactly.xc_participant) );''')

        common.append('''create step if not exists s_order_validate_duplicates as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, '' as error_field, 'Duplicate orders' AS reject_reason from delta.prestage_order_item group by order_code,item_code having count(*)> 1);''')

        common.append('''create step if not exists s_order_validate_missing_assignments as (insert into Delta(TableName='delta.process_log') select CurDate() AS process_log_creation_date,:v_period_name AS processing_period,'REJECT' as category,order_code, item_code, '' as error_field, 'Assignment missing' AS reject_reason from delta.prestage_order_item where order_code||item_code not inn (select distinct order_code||item_code from delta.prestage_order_item_assignment where order_code||item_code not inn (select distinct order_code||item_code from  delta.process_log where category='REJECT')));''')

        common.append('''create step if not exists s_shared_create_batches as (incent create batches);''')

        common.append('''create step if not exists s_shared_insert_staging_customer as (insert into staging.customer (action, name) select distinct 'save', customer_name from staging.order_item where customer_name not in (select name from xactly.xc_customer) and customer_name is not null);''')

        common.append('''create step if not exists s_shared_upload_customers as (incent upload customers);''')

        common.append('''create step if not exists s_shared_insert_staging_product as (insert into staging.product (action, name) select distinct 'save', product_name from staging.order_item where product_name not in (select name from xactly.xc_product) and product_name is not null);''')

        common.append('''create step if not exists s_shared_upload_products as (incent upload products);''')

        common.append('''create step if not exists s_shared_insert_staging_geography as (insert into staging.geography (action, name) select distinct 'save', geography_name from staging.order_item where geography_name not in (select name from xactly.xc_geography) and geography_name is not null);''')

        common.append('''create step if not exists s_shared_upload_geographies as (incent upload geographies);''')

        common.append('''create step if not exists s_shared_validate_orders as (incent validate orders);''')

        common.append('''create step if not exists s_archive_invalid_order_errors as (insert into Delta(TableName='delta.archive_order_item_validation_error', Overwrite=true , Unlogged=true) select * from staging.order_item_validation_error);''')

        common.append('''create step if not exists s_copy_valid_stage_order_item_temp as (insert into Delta(TableName='delta.valid_stage_order_item_temp', Overwrite=true , Unlogged=true) select * from staging.order_item where order_code||item_code not inn (select distinct order_code||item_code from staging.order_item_validation_error ));''')

        common.append('''create step if not exists s_copy_valid_stage_order_item_assignment_temp as (insert into Delta(TableName='delta.valid_stage_order_item_assignment_temp', Overwrite=true , Unlogged=true) select *  from staging.order_item_assignment where order_code||item_code not inn (select distinct order_code||item_code from staging.order_item_validation_error ));''')

        common.append('''create step if not exists s_shared_upload_orders as (incent upload orders (Validate=false));''')

        common.append('''create step if not exists s_set_sleep_60 as (sleep 60);''')

        common.append('''create step if not exists s_set_sleep_30 as (sleep 30);''')

        common.append('''create step if not exists s_create_i_insert_staging_orders_by_batch as (create iterator if not exists i_insert_staging_orders_by_batch for step s_insert_into_staging_with_seq  over select distinct batch_name as v_batch_name from delta.valid_stage_order_item_temp);''')

        common.append('''create step if not exists s_invoke_i_insert_staging_orders_by_batch as (invoke iterator i_insert_staging_orders_by_batch);''')

        common.append('''create step if not exists s_set_order_item_field_list as (set v_order_item_field_list *= select GatherString(name,',') from (describe select * from staging.order_item) where name <>'batch_name');''')

        common.append('''create step if not exists s_insert_into_delta_valid_order_assignments_unique as (Insert into Delta(TableName='delta.tmp_unique_order_assignments', Unlogged=true, Overwrite=true,PrimaryKey='order_code,item_code') select distinct order_code,item_code from delta.valid_stage_order_item_assignment_temp);''')

        common.append('''create step if not exists s_insert_into_staging_with_seq as (insert into staging.order_item  ({:v_order_item_field_list},batch_name) select {:v_order_item_field_list},SubString(batch_name,0,StrLength(batch_name)-3)||FormatNumber(((SeqNum()-1)/:v_batch_size)+1,'000') as new_batch_name from  delta.valid_stage_order_item_temp a join  delta.tmp_unique_order_assignments b on a.order_code=b.order_code and a.item_code=b.item_code where a.batch_name=:v_batch_name);''')

        common.append('''create step if not exists s_insert_into_staging_order_assignments as (insert into staging.order_item_assignment (order_code, item_code, employee_id, split_amount_pct) select order_code,item_code,employee_id,split_amount_pct from delta.valid_stage_order_item_assignment_temp);''')

        common.append('''create step if not exists s_create_prestage_order_item as (Insert into Delta(TableName='delta.prestage_order_item', Unlogged=true, Overwrite=true)  select * from staging.order_item where '1'='2');''')

        common.append('''create step if not exists s_create_prestage_order_item_assignment as (Insert into Delta(TableName='delta.prestage_order_item_assignment', Unlogged=true, Overwrite=true) select * from staging.order_item_assignment where '1'='2');''')

        common.append('''create step if not exists s_create_temp_dir as (call MakeDir(Directory= '/logs/temp') );''')

        common.append('''create step if not exists s_delete_temp_dir as (call DeleteFile(Directory='/logs/temp' , Recursive=true) );''')

        common.append('''create step if not exists s_shared_set_process_id as ( set v_process_id *= select id from (show invocations) where object_name like :v_pipeline_name order by created_instant desc );''')

        common.append('''create step if not exists s_shared_create_invocation_details_log as ( Insert into Delta(TableName='delta.invocation_details_log', Overwrite=true,Unlogged=true) select object_id, parent_name, object_name, FormatDateTime(WithZoneSameInstant(created_instant, 'EST5EDT'), 'yyyy-MM-dd HH:mm:ss z') as EST_start_time, FormatDateTime(WithZoneSameInstant(completed_instant, 'EST5EDT'), 'yyyy-MM-dd HH:mm:ss z') as EST_end_time, DiffTime(created_instant,completed_instant,'SECONDS')/60.0 as processing_time_minutes,status,exception_message, reason_code, command from (show invocation details ) where invocation_id=:v_process_id and object_type in ('step','pipeline') order by created_instant,parent_id );''')

        common.append('''create step if not exists s_shared_create_invocation_details_log_file as ( call  WriteFile(FilePath= '/logs/invocation_details_'||:v_process_id||'.csv', FirstLineNames=true, Input=(select * from delta.invocation_details_log ), Separator=',', Quote='"') );''')

        common.append('''create step if not exists s_shared_set_process_status as (set v_process_exec_status *= select Uppercase(status) from  delta.invocation_details_log where object_name like :v_pipeline_name);''')

        common.append('''create step if not exists s_shared_create_e_process_error as (create email if not exists e_process_error as ("From"=:v_email_from, "To"=:v_email_to, "Body"=:v_generic_email_body, "BodyType"='html', "Subject"=:v_shared_customer_name ||' - Invocation logs - pipeline  "' ||:v_pipeline_name|| '" has completed with status - "'|| :v_process_exec_status ||'" for period ' || :v_period_name , "Attachment1" = ('/logs/invocation_details_' ||:v_process_id||'.csv'),"AttachmentType1" = 'plain' ));''')

        common.append('''create step if not exists s_shared_send_e_process_error as (send email e_process_error);''')

        common.append('''create step if not exists s_create_delta_pglist_with_calenderyear as (Insert into Delta(TableName='delta.pglist_with_calenderyear',Overwrite=true, Unlogged=true) select pgs.child_process_group_name as process_group_name, pgs.period_name, min(pgs.process_group_order) as process_group_order , p2.period_id as plan_year from delta.pgqueue_mapping_with_period pgs join xactly.xc_period p1 on p1.name =pgs.period_name join xactly.xc_period p2 on LookupPeriodTypeById(p2.period_type_id_fk) ='YEARLY' where p1.start_date between p2.start_date and p2.end_date  group by child_process_group_name,period_name,plan_year);''')

        common.append('''create step if not exists s_shared_pg_create_xactly_proc_group_tables as (create table if not exists xactly.xc_proc_group using Xactly(TableName='xc_proc_group'); create table if not exists xactly.xc_proc_group_version using Xactly(TableName='xc_proc_group_version'); create table if not exists xactly.xc_proc_group_step using Xactly(TableName='xc_proc_group_step'); create table if not exists xactly.xc_proc_group_step_param using Xactly(TableName='xc_proc_group_step_param'); create table if not exists xactly.xc_ext_proc using Xactly(TableName='xc_ext_proc'); create table if not exists xactly.xc_ext_proc_param using Xactly(TableName='xc_ext_proc_param'););''')

        common.append('''create step if not exists s_shared_pg_load_incent_pg_data as (Insert into Delta(TableName='delta.pg_dump',Unlogged=true, Overwrite=true)  select distinct plan_year, a.name as process_group_name, c.name as step_name,e.name as ext_proc_name,f.name param_name,Uppercase(f.name) ucase_param_name,f.value  as param_default_value ,c.priority  as step_order  from xactly.xc_proc_group a join xactly.xc_proc_group_version b on  a.proc_group_id=b.proc_group_id   left join  xactly.xc_proc_group_step c on  c.proc_group_version_id=b.proc_group_version_id and action_type='External_Process'  left join xactly.xc_proc_group_step_param d on  d.proc_group_step_id=c.proc_group_step_id and d.name='ExternalProcess'  left join xactly.xc_EXT_PROC e on e.EXT_PROC_ID=d.VALUE_STRING left join xactly.xc_ext_proc_param f on e.EXT_PROC_ID=f.EXT_PROC_ID order by a.name,c.priority);''')

        common.append('''create step if not exists s_shared_pg_get_final_pgs_to_queue as (Insert into Delta(TableName='delta.final_pg_list_to_queue',Unlogged=true, Overwrite=true)  select  pgd.process_group_name,pgy.period_name,  pgd.step_name, pgd.ext_proc_name, pgd.param_name , if (pgm.param_name is not null) then if (pgm.is_variable=true) then Replace(vars."value",'''','') else pgm.param_value end else pgd.param_default_value end as param_value ,pgy.process_group_order,pgd.step_order  from delta.pg_dump pgd join  delta.pglist_with_calenderyear pgy on pgd.process_group_name=pgy.process_group_name and pgd.plan_year=pgy.plan_year left join delta.pgqueue_mapping_with_period pgm on pgd.process_group_name=pgm.child_process_group_name and pgd.ucase_param_name=pgm.param_name  left join  (show variables) vars on pgm.param_value=vars.name order by pgy.process_group_order, pgd.step_order );''')

        common.append('''create step if not exists s_shared_pg_generate_param_json as (Insert into Delta(TableName='delta.pg_to_queue_paramlist',Unlogged=true, Overwrite=true)  select process_group_name,period_name,step_name,step_order,process_group_order,'['||GatherString('{"'||param_name||'":"'||param_value||'"}',',')||']' as param_list from delta.final_pg_list_to_queue group by process_group_name,period_name,step_name ,step_order,process_group_order order by process_group_order, step_order );''')

        common.append('''create step if not exists s_shared_pg_generate_pg_json as (Insert into Delta(TableName='delta.pg_to_queue_withparam_json',Overwrite=True,Unlogged=true) select process_group_name,period_name,process_group_order,'['||GatherString('{"'||step_name||'":'||param_list||'}',',')||']' as parameter_overrides from delta.pg_to_queue_paramlist group by process_group_name,period_name,process_group_order order by process_group_order);''')

        common.append('''create step if not exists s_shared_incent_queue_pg as (Insert into Delta(TableName='delta.queuepg_status_log',Overwrite=true,Unlogged=true) select  CurDateTime() as created_instant,:v_process_name as parent_process_group ,process_group_name,period_name,parameter_overrides,status_code,status_message,exception_message,queue_region_id from (call QueueIncentProcessGroup(Input=(select process_group_name, period_name,  Nvl(parameter_overrides,'[]') as parameter_overrides from delta.pg_to_queue_withparam_json ))) t1 );''')

        common.append('''create step if not exists s_set_queuepg_log_file_name as (set v_queuepg_log_file_name *= Replace(:v_process_name,' ','_')||'_' ||:v_period_name ||'_QueuePG_Log_' || FormatDateTime(CurDateTime(),'yyyyMMdd-HH.mm.ss') ||'.csv');''')

        common.append('''create step if not exists s_create_queuepg_log_file as (call  WriteFile(FilePath='/logs/' ||:v_queuepg_log_file_name,Input= (select * from delta.queuepg_status_log), FirstLineNames=true,Separator=',',Quote='"',Trim=true));''')

        common.append('''create step if not exists s_onetime_create_pgconfig_schema as (create schema  if not exists pgconfig);''')

        common.append('''create step if not exists s_onetime_create_table_pgqueue_mapping as (Insert into Delta(TableName='pgconfig.pgqueue_mapping',Overwrite=false, Unlogged=false) select '' as parent_process_group_name, '' as child_process_group_name, 0 as process_group_order ,'' as param_name,  '' as param_value , false as is_variable from Empty(rows=0) );''')

        common.append('''create step if not exists s_onetime_setup_add_pgs as (Insert into Delta(TableName='pgconfig.pgqueue_mapping') select  parent_process_group_name, child_process_group_name, process_group_order, param_name, param_value ,  is_variable  from ReadFile(FilePath='/pgconfig/Connect_PGSetup.csv' , FirstLineNames=true, Separator=',', Quote='"', Trim=true) where parent_process_group_name||child_process_group_name||process_group_order||param_name not inn (select distinct parent_process_group_name||child_process_group_name||process_group_order||param_name from pgconfig.pgqueue_mapping ));''')



        common.append('''--start: common shared pipeline to set the variables ''') 

        common.append('''create pipeline if not exists p_set_dynamic_variables (ContinueOnError=false);''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_period_start_date;''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_period_end_date;''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_set_v_email_to;''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_set_generic_email_body;''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_set_generic_email_subject;''')


        common.append('''alter pipeline p_set_dynamic_variables add step s_create_generic_email;''')

        common.append('''--end :  common shared pipeline to set the variables  ''')




        common.append('''--start: common shared pipeline to queue pgs from the table delta.pgqueue_mapping_with_period   ''')

        common.append('''create pipeline if not exists p_shared_queue_pg (ContinueOnError=false);''')


        common.append('''alter pipeline p_shared_queue_pg add step s_create_delta_pglist_with_calenderyear;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_pg_create_xactly_proc_group_tables;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_pg_load_incent_pg_data;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_pg_get_final_pgs_to_queue;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_pg_generate_param_json;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_pg_generate_pg_json;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_shared_incent_queue_pg;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_set_queuepg_log_file_name;''')


        common.append('''alter pipeline p_shared_queue_pg add step s_create_queuepg_log_file;''')

        common.append('''--end :  common shared pipeline to queue pgs from the table delta.pgqueue_mapping_with_period ''')  




        common.append('''--start: pipeline to setup pgs ( needs to be run only once when setting up for the first time)''')   

        common.append('''create pipeline if not exists p_onetime_setup_pgqueue_mapping (ContinueOnError=false);''')


        common.append('''alter pipeline p_onetime_setup_pgqueue_mapping add step s_onetime_create_pgconfig_schema;''')


        common.append('''alter pipeline p_onetime_setup_pgqueue_mapping add step s_onetime_create_table_pgqueue_mapping as (Condition=select if count(*)=0   then true else false end from (show tables) where schema_name='pgconfig' and name ='pgqueue_mapping', OnConditionFalse=s_shared_no_action, AbortOnConditionFalse=false);''')


        common.append('''alter pipeline p_onetime_setup_pgqueue_mapping add step s_onetime_setup_add_pgs;''')

        common.append('''--end :  pipeline to setup pgs ( needs to be run only once when setting up for the first time)  ''') 




        common.append('''--start: common shared pipeline to send the validations errors to email id set in variable v_email_to_error. Please make sure to set this variable before invoking this pipeline''')  

        common.append('''create pipeline if not exists p_email_invocation_errors (ContinueOnError=false);''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_set_process_id;''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_create_invocation_details_log;''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_set_process_status;''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_create_invocation_details_log_file;''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_create_e_process_error;''')


        common.append('''alter pipeline p_email_invocation_errors add step s_shared_send_e_process_error;''')

        common.append('''--end :  common shared pipeline to send the validations errors to email id set in variable v_email_to_error. Please make sure to set this variable before invoking this pipeline''')  




        common.append('''--start: common shared pipeline to clear all staging tables  ''')

        common.append('''create pipeline if not exists p_shared_delete_staging_tables (ContinueOnError=false);''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_stage_order_item;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_stage_order_item_asgnmt;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_stage_order_item_val_err;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_geography;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_geography_exception;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_customer;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_customer_exception;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_product;''')


        common.append('''alter pipeline p_shared_delete_staging_tables add step s_delete_staging_product_exception;''')

        common.append('''--end :  common shared pipeline to clear all staging tables''')  




        common.append('''--start: common shared pipeline to perform stadard order field validations ''') 

        common.append('''create pipeline if not exists p_standard_order_validations (ContinueOnError=false);''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_mandatory_fields;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_batch_type;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_amount_unit_type;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_order_type;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_employee_id;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_duplicates;''')


        common.append('''alter pipeline p_standard_order_validations add step s_order_validate_missing_assignments;''')

        common.append('''--end :  common shared pipeline to perform stadard order field validations''')  



        common.append('''create pipeline if not exists p_shared_upload_orders (ContinueOnError=false);''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_create_batches;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_insert_staging_customer;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_upload_customers;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_insert_staging_product;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_upload_products;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_insert_staging_geography;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_upload_geographies;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_set_sleep_60;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_validate_orders;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_archive_invalid_order_errors;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_copy_valid_stage_order_item_temp;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_copy_valid_stage_order_item_assignment_temp;''')


        common.append('''alter pipeline p_shared_upload_orders add pipeline p_shared_delete_staging_tables;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_insert_into_delta_valid_order_assignments_unique;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_insert_into_staging_order_assignments;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_create_i_insert_staging_orders_by_batch;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_invoke_i_insert_staging_orders_by_batch;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_create_batches;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_set_sleep_30;''')


        common.append('''alter pipeline p_shared_upload_orders add step s_shared_upload_orders;''')

        print('Common steps completed')
        #
        np.savetxt(f'''{process_short_name}_connect_common_steps.txt''',common, fmt='%s', delimiter="\t")
        files.append(f'''{process_short_name}_connect_common_steps.txt''')
        ###return send_file(f'''{process_short_name}_connect_common_steps.txt''',as_attachment=True)
        
        with zipfile.ZipFile(UPLOAD_PATH + 'Test.zip', 'w') as zipF:
            for file in files:
                zipF.write(UPLOAD_PATH + file, basename(UPLOAD_PATH + file), compress_type=zipfile.ZIP_DEFLATED)
                os.remove(UPLOAD_PATH+file)
            zipF.close()
            return
       

        #response = Response(np.savetxt(f'''{process_short_name}_connect_common_steps.txt''',common, fmt='%s', delimiter="\t"), mimetype='text/csv')
        # add a filename
        #response.headers.set("Content-Disposition", "attachment", filename="log.txt")
        #return response
        return render_template('data.html')
    
       

if __name__=='__main__':
    app.debug = True
    app.run()



# url = "http://url.to.file/sample.docx"
# docx = BytesIO(requests.get(url).content)

# # extract text
# text = docx2txt.process(docx)